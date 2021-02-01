# -*- coding: utf-8 -*-
"""
Created on Mon Aug  3 11:23:41 2020

@author: cheng
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.shared import Inches
from datetime import datetime
import pandas as pd
from jqdatasdk import finance
from jqdatasdk import *

auth('13739188902','ZNnb20160801')
def write_docx(filename, sd,risk,duration,fig1,p_per1, p_money1,fig2,p_per2, p_money2,md2,piename,w):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    H=document.add_heading("",level=0)
    run=H.add_run(filename + '理财与基金对比报告'+'('+risk+')')
    run.font.name=u'Cambria'
    run.font.size=Pt(18)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    today = datetime.today()
    H=document.add_heading("",level=1)
    run=H.add_run(today.date().__str__())
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    H=document.add_heading("",level=1)
    run=H.add_run('一、真实历史交易复现：')
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    p = document.add_paragraph('复现周期：')
    p.add_run(sd.date().__str__() + '——' + today.date().__str__())
    document.add_picture(fig1, width=Inches(6.0))
    p=document.add_paragraph('')                          #zym 
    p=document.add_paragraph('期间收益：'+str(round(p_per1*100,2))+'%('+str(round(p_money1/10000))+'万元)')
    p=document.add_paragraph('')
    
    H=document.add_heading("",level=1)
    run=H.add_run('二、'+risk+'组合公募基金投资对比复现')
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    p=document.add_paragraph('持仓组合类别：'+risk+'组合')
    document.add_paragraph('投资方法：每次申购均买入对应季度的'+risk+'组合，买入后每持有'+duration+'调整一次')
    p = document.add_paragraph('复现周期：')
    p.add_run(sd.date().__str__() + '——' + today.date().__str__())
    document.add_picture(fig2, width=Inches(6.0))
    document.add_paragraph('')                                   #zym 
    p=document.add_paragraph('期间收益：'+str(round(p_per2*100,2))+'%('+str(round(p_money2/10000))+'万元)')
    p.add_run(change(p_money1, p_money2))
    document.add_paragraph('组合最大回撤（可能面临的最大风险）：'+str(round(md2*100,2))+'%')
    document.add_paragraph('')
    
    H=document.add_heading("",level=1)
    run=H.add_run('三、推荐基金组合')
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    document.add_paragraph('组合类别：'+risk+'公募基金组合')
    year = today.year
    quarter = (today.month - 1) // 3 + 1
    document.add_paragraph('组合更新时间：'+str(year)+'年'+str(quarter)+'季度（组合明细如下图所示）')
    document.add_picture(piename, width=Inches(6.0))
    document.add_paragraph('')
    document.add_paragraph('组合明细（以100万元计算）：')
    length=len(w['weights'])
    # mw=w.values
    for i in range(0,length):
        #print(mw[i][0])
        #print(mw[i][1])
        name=finance.run_query(query(finance.FUND_MAIN_INFO).filter(finance.FUND_MAIN_INFO.main_code==w.index[i]))['name']
        stri=w.index[i]+name[0]+'——'+str(round(float(w.iloc[i][0])*100))+'万'
        document.add_paragraph(stri)
    
    document.save(os.path.dirname(__file__)+'\\result\\'+filename + '理财与基金对比报告'+'('+risk+').docx')
    

def change(last, now):
    if now<0 and last<0:
        d=round((now / last - 1) * 100)
        return '亏损比原来减少了'+str(-d)+'%'
    elif now<0 and last>0:
        return ''
    elif now > last :
        if last < 0 and now > 0:
            d = round((now / -last) * 100)
        elif last * now > 0:
            d = round((now / last - 1) * 100)
        return '比原来提高了'+str(d)+'%'
    else:
        return ''
    

    
    

# w=pd.DataFrame({'':['485111','519736','000771'],'weights':[0.7654,0.2283,0.009]})    
# write_docx('客户A', datetime(2015,5,18),'低风险','1年',w)

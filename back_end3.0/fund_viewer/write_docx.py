# -*- coding: utf-8 -*-
"""
Created on Mon Aug  3 11:23:41 2020

@author: cheng
"""


from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.shared import Inches
from datetime import datetime
import pandas as pd
from jqdatasdk import finance
from jqdatasdk import *
import os

auth('13739188902', 'Zn20160801')
def write_docx(filename, sd,risk,duration,fig1,result1, p_money1,md_money1,fig2,result2, p_money2,md_money2,piename,w):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    H=document.add_heading("",level=0)
    run=H.add_run(filename + '基金检视报告（'+risk+'版）')
    run.font.name=u'Cambria'
    run.font.size=Pt(18)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    today = datetime.today()
    H=document.add_heading("",level=1)
    run=H.add_run(today.date().__str__())
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    H=document.add_heading("",level=1)
    run=H.add_run('一、真实历史交易复现：')
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    p = document.add_paragraph('复现周期：')
    p.add_run(sd.date().__str__() + '——' + today.date().__str__())
    document.add_picture(fig1, width=Inches(6.0))
    p=document.add_paragraph('')
    p=document.add_paragraph('期间收益：'+str(round(result1['r']*100,2))+'%('+str(round(p_money1/10000))+'万元)')
    document.add_paragraph('期间最大回撤：'+str(round(result1['md']*100,2))+'%('+str(round(md_money1/10000))+'万元)')
    document.add_paragraph('期间风险收益比：'+str(round(result1['r/md'],2))+'（即总收益/最大回撤）')
    p=document.add_paragraph('')
    
    H=document.add_heading("",level=1)
    run=H.add_run('二、组合化投资对比复现')
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    p=document.add_paragraph('持仓组合类别：'+risk+'组合')
    document.add_paragraph('投资方法：每次申购均买入对应季度的'+risk+'组合，买入后每持有'+duration+'个月调整一次')
    p = document.add_paragraph('复现周期：')
    p.add_run(sd.date().__str__() + '——' + today.date().__str__())
    document.add_picture(fig2, width=Inches(6.0))
    document.add_paragraph('')
    p=document.add_paragraph('期间收益：'+str(round(result2['r']*100,2))+'%('+str(round(p_money2/10000))+'万元)')
    p.add_run(change(p_money1, p_money2))
    p=document.add_paragraph('期间最大回撤：'+str(round(result2['md']*100,2))+'%('+str(round(md_money2/10000))+'万元)')
    p.add_run(changemd(result1['md'], result2['md']))
    p=document.add_paragraph('期间风险收益比：'+str(round(result2['r/md'],2))+'（即总收益/最大回撤）')
    p.add_run(changermd(result1['r/md'], result2['r/md']))
    document.add_paragraph('')
    
    H=document.add_heading("",level=1)
    run=H.add_run('三、推荐基金组合')
    run.font.name=u'Cambria'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0,0,0)
    
    document.add_paragraph('组合类别：'+risk+'公募基金组合')
    year = today.year
    quarter = (today.month - 1) // 3 + 1
    document.add_paragraph('组合更新时间：'+str(year)+'年'+str(quarter)+'季度（组合明细如下图所示）')
    document.add_picture(piename, width=Inches(6.0))
    document.add_paragraph('')
    document.add_paragraph('组合明细（以100万元计算）：')
    length=len(w['weights'])
    # mw=w.values
    for i in range(0,length):
        #print(mw[i][0])
        #print(mw[i][1])
        name=finance.run_query(query(finance.FUND_MAIN_INFO).filter(finance.FUND_MAIN_INFO.main_code==w.index[i]))['name']
        stri=w.index[i]+name[0]+'——'+str(round(float(w.iloc[i][0])*100))+'万'
        #stri=w.index[i]+'——'+str(round(float(w.iloc[i][0])*100))+'万'
        document.add_paragraph(stri)
    
    document.save(os.path.dirname(__file__)+'\\'+'result\\'+filename + '基金检视报告（'+risk+'版）.docx')
    

def change(last, now):
    if now<0 and last<0:
        d=round((now / last - 1) * 100)
        return '亏损比原来减少了'+str(-d)+'%'
    elif now<0 and last>0:
        return ''
    elif now > last :
        if last < 0 and now > 0:
            d = round((now / -last) * 100)
        elif last * now > 0:
            d = round((now / last - 1) * 100)
        return '比原来提高了'+str(d)+'%'
    else:
        return ''
    
def changemd(last, now):
    if now < last:
        d = round((1- now / last) * 100)
        return '比原来降低了'+str(d)+'%'
    else:
        return ''

def changermd(last, now):
    if now > last :
        if last<0 and now<0:
            d=round(((-now+last)*100/last)+0.00001)
        elif last < 0 and now > 0:
            d = round((((now-last) / -last) * 100)+0.00001)
        elif last * now > 0:
            d = round(((now / last - 1) * 100)+0.00001)
        return '比原来提高了'+str(d)+'%'
    else:
        return '' 

    
    

# w=pd.DataFrame({'':['485111','519736','000771'],'weights':[0.7654,0.2283,0.009]})    
# write_docx('客户A', datetime(2015,5,18),'低风险','1年',w)
